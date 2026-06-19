"""Document parser — extracts plain text from various file formats."""

import os


def parse_document(file_path: str, file_type: str) -> str:
    """Parse a document file and return its text content."""
    parsers = {
        ".txt": _parse_txt,
        ".md": _parse_md,
        ".pdf": _parse_pdf,
        ".docx": _parse_docx,
    }
    ext = file_type.lower()
    parser = parsers.get(ext)
    if parser is None:
        raise ValueError(f"Unsupported file type: {ext}")
    return parser(file_path)


def _parse_txt(path: str) -> str:
    """Parse a plain text file."""
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _parse_md(path: str) -> str:
    """Parse a markdown file (preserve structure)."""
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _parse_pdf(path: str) -> str:
    """Parse a PDF file using PyMuPDF with enhanced extraction."""
    import fitz  # pymupdf

    doc = fitz.open(path)
    pages = []
    total_text_len = 0

    for i, page in enumerate(doc):
        blocks = page.get_text("dict")["blocks"]
        page_parts = []

        for block in blocks:
            # Text block
            if block["type"] == 0:
                for line in block.get("lines", []):
                    text = "".join(span["text"] for span in line.get("spans", []))
                    text = text.strip()
                    if text:
                        page_parts.append(text)

            # Table block (detect via line art or grouped rects)
            if block["type"] == 1:
                # Image — skip during normal text extraction
                continue

        text = "\n".join(page_parts) if page_parts else ""
        if text.strip():
            total_text_len += len(text)
            pages.append(f"--- 第 {i + 1} 页 ---\n{text}")
        else:
            pages.append(f"--- 第 {i + 1} 页 ---\n(空白页)")

    doc.close()

    result = "\n\n".join(pages)

    # If very little text was extracted (< 100 chars per page), it's likely scanned/image PDF.
    # In that case, the image-based document parsing (multimodal LLM) should be used instead.
    return result


def _parse_docx(path: str) -> str:
    """Parse a DOCX file using python-docx (paragraphs + tables)."""
    from docx import Document

    doc = Document(path)
    parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract table content (key content often lives in tables)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(cells)
            if row_text.strip():
                rows.append(row_text)
        if rows:
            parts.append("[表格]\n" + "\n".join(rows))

    return "\n\n".join(parts)
